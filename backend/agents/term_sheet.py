"""Term Sheet Generator — produces a professional deal term sheet DOCX.

Modelled on the UAE Ed71 deal term sheet format.
This is one of the two primary deliverables to the country/state
(the other being the investor deck).

Includes a HITL gate for term sheet assumptions (deal terms)
that the user can review/adjust before the term sheet is generated.
Changes to key financial terms trigger financial model recalculation.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt as DocxPt, RGBColor as DocxRGB, Inches as DocxInches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from models.schemas import (
    CountryProfile, EducationAnalysis, Strategy,
    FinancialModel, FinancialAssumptions, FinancialAssumption, TargetType,
)
from services.llm import call_llm_plain
from config import OUTPUT_DIR

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# UAE Heads of Terms — reference template
# ---------------------------------------------------------------------------

UAE_HOT_REFERENCE = """
## REFERENCE: UAE Ed71 Heads of Terms (ASM / Alpha Holdings × UAE)

This is the actual deal structure used for the UAE partnership. Use it as a
template and calibrate all new term sheets to this standard.

### 1. Strategic Investments (optional, illustrative)
| Investment        | Amount | Valuation | Stake | Timing |
|-------------------|--------|-----------|-------|--------|
| Incept Labs       | $1B    | $10B      | 10%   | 2026   |
| Alpha Holdings    | $100M  | $10B      | 1%    | 2026   |
| Alpha Holdings    | $1B    | $100B     | 1%    | 2027   |
| TimeBack EdGame   | $20M   | $1B       | 2%    | 2025   |

### 2a. Alpha Schools — Original Scope
2 Flagships ($100k tuition, 2k students) + 20 Branded ($50k tuition, 20k students) = 22,000 students.

| Item                          | Up-Front | Ongoing  | Recipient      |
|-------------------------------|----------|----------|----------------|
| Launch Fee                    | $250M    |          | Alpha Holdings |
| Scholarship/Tuition Backstop  |          | $350M/yr | Alpha Holdings |
| **Total**                     | **$250M**| **$1.75B** (5yr) |        |

### 2b. Alpha Schools — Reduced Scope (flagship only)
2 Flagships ($100k), 2,000 students. 50% fill backstop for 5 years, prepaid.

| Item                          | Up-Front | Ongoing | Recipient      |
|-------------------------------|----------|---------|----------------|
| Scholarship/Tuition Backstop  | $500M    |         | Alpha Holdings |

### 3. Ed71 JV Schools
100% Next71 owned, Alpha exclusive operator.
$25k annual tuition × 100-200k students = $2.5B-$5.0B annual revenue.

| Item                          | Up-Front | Ongoing     | Recipient      | Notes |
|-------------------------------|----------|-------------|----------------|-------|
| Incept UAE eduLLM Dev         | $250M    |             | Alpha Holdings | Optional |
| AlphaCore License             | $250M    |             | Alpha Holdings | |
| UAE EdTech App R&D            | $250M    |             | Local expense  | |
| AsasCore R&D                  | $250M    |             | Local expense  | |
| Marketing/Launch Capital      |          | $250M       | Local expense  | |
| Scholarships                  |          | $1B/100k    | Local expense  | |
| Management Fee                | $250M    | 10% tuition | Alpha Holdings | Min $2,500/student |
| TimeBack Fee                  |          | 20% tuition | Alpha Holdings | Min $5,000/student |
| **Total Ed71 Schools**        | **$1.25B** | **Scale dependent** | | |

### Key Principles
- Management fee: 10% of tuition (non-negotiable floor), $2,500 per student minimum
- Timeback license: 20% of tuition (non-negotiable floor), $5,000 per student minimum
- Upfront prepayment: 2 years of management fees at signing
- Flagship school presence essential for "halo effect"
- Scholarship backstop: 50% capacity × tuition × number of schools × 5 years
- Alpha commitments: (1) children love school, (2) learn 2x faster, (3) life skills for AI age
- Cultural IP layer: local identity, language, and values fully integrated
"""

# ---------------------------------------------------------------------------
# Term Sheet Assumptions Generator
# ---------------------------------------------------------------------------

# Keys that overlap with the financial model and require recalculation
FINANCIAL_MODEL_OVERLAP_KEYS = {
    "ts_upfront_ip_fee",
    "ts_management_fee_pct",
    "ts_timeback_license_pct",
    "ts_students_year5",
    "ts_per_student_budget",
    "ts_capex_per_school",
}

# Mapping from term sheet assumption keys to financial model assumption keys
TS_TO_FM_KEY_MAP = {
    "ts_upfront_ip_fee": "upfront_ip_fee",
    "ts_management_fee_pct": "management_fee_pct",
    "ts_timeback_license_pct": "timeback_license_pct",
    "ts_students_year5": "students_year5",
    "ts_per_student_budget": "per_student_budget",
    "ts_capex_per_school": "capex_per_school",
}


def generate_term_sheet_assumptions(
    target: str,
    country_profile: CountryProfile,
    strategy: Strategy,
    financial_model: FinancialModel,
    financial_assumptions: FinancialAssumptions,
) -> FinancialAssumptions:
    """Generate a set of configurable term sheet assumptions for HITL review.

    These are the deal-specific terms that will appear in the term sheet.
    Some overlap with the financial model — if user changes those, a
    recalculation of the financial model is required.
    """
    logger.info("Generating term sheet assumptions for %s", target)

    # Pull current values from financial model/assumptions
    fa = {item.key: item.value for item in financial_assumptions.assumptions}

    mgmt_pct = round(financial_model.management_fee_pct * 100)
    timeback_pct = round(financial_model.timeback_license_pct * 100)
    upfront_ip = financial_model.upfront_ip_fee
    upfront_ip_m = round(upfront_ip / 1_000_000) if upfront_ip > 1_000_000 else upfront_ip

    # Student targets from financial model
    y5_students = 0
    if financial_model.pnl_projection:
        y5_students = financial_model.pnl_projection[-1].students

    y5_revenue = 0
    if financial_model.pnl_projection:
        y5_revenue = financial_model.pnl_projection[-1].revenue

    per_student = fa.get("per_student_budget", fa.get("mid_tuition", 20_000))
    capex_per_school = fa.get("capex_per_school", 5_000_000)

    # Calculate backstop based on UAE formula: 50% capacity × tuition × schools × 5 years
    num_schools = financial_model.pnl_projection[-1].schools if financial_model.pnl_projection else 10
    backstop_annual = round(0.5 * y5_students * per_student / 1_000_000)  # $M
    mgmt_fee_prepay_m = round(min(5000, y5_students) * per_student * (mgmt_pct / 100) * 2 / 1_000_000)

    is_us_state = country_profile.target.type == TargetType.US_STATE

    assumptions = [
        # --- Deal Structure ---
        FinancialAssumption(
            key="ts_partnership_type",
            label="Partnership Structure",
            value=1,  # 1=JV, 2=Licensing, 3=Direct
            min_val=1, max_val=3, step=1,
            unit="type", category="deal",
            description="1 = Joint Venture, 2 = Licensing, 3 = Direct Operation",
        ),
        FinancialAssumption(
            key="ts_alpha_ownership_pct",
            label="Alpha Ownership / Control (%)",
            value=49 if not is_us_state else 100,
            min_val=10, max_val=100, step=1,
            unit="%", category="deal",
            description="Alpha's equity stake in the JV (100% for direct)",
        ),
        FinancialAssumption(
            key="ts_term_years",
            label="Partnership Term (years)",
            value=25, min_val=10, max_val=50, step=5,
            unit="years", category="deal",
            description="Initial partnership duration (renewable)",
        ),
        FinancialAssumption(
            key="ts_exclusivity_years",
            label="Exclusivity Period (years)",
            value=25, min_val=5, max_val=50, step=5,
            unit="years", category="deal",
            description="Alpha has exclusive rights to the 2hr Learning model in the territory",
        ),

        # --- Fee Structure (overlaps with financial model) ---
        FinancialAssumption(
            key="ts_upfront_ip_fee",
            label="Upfront IP / Development Fee ($M)",
            value=upfront_ip_m,
            min_val=5, max_val=500, step=5,
            unit="$M", category="fees",
            description="One-time IP licensing and development fee. ⚠️ Changes recalculate financial model.",
        ),
        FinancialAssumption(
            key="ts_management_fee_pct",
            label="Management Fee (% of school revenue)",
            value=mgmt_pct,
            min_val=5, max_val=20, step=1,
            unit="%", category="fees",
            description="Alpha's ongoing management fee. 10% floor is non-negotiable. ⚠️ Changes recalculate financial model.",
            locked=True,
        ),
        FinancialAssumption(
            key="ts_timeback_license_pct",
            label="Timeback License (% of per-student budget)",
            value=timeback_pct,
            min_val=10, max_val=30, step=1,
            unit="%", category="fees",
            description="Alpha's technology/IP licensing fee. 20% floor is non-negotiable. ⚠️ Changes recalculate financial model.",
            locked=True,
        ),
        FinancialAssumption(
            key="ts_mgmt_fee_prepay_years",
            label="Management Fee Prepayment (years)",
            value=2, min_val=0, max_val=5, step=1,
            unit="years", category="fees",
            description="Years of management fee prepaid at signing",
        ),
        FinancialAssumption(
            key="ts_mgmt_fee_prepay_amount",
            label="Management Fee Prepayment Amount ($M)",
            value=max(1, mgmt_fee_prepay_m),
            min_val=0, max_val=500, step=5,
            unit="$M", category="fees",
            description="Calculated: first-year students × per-student budget × mgmt fee % × prepay years",
        ),

        # --- School Portfolio (overlaps with financial model) ---
        FinancialAssumption(
            key="ts_students_year5",
            label="Year 5 Student Target",
            value=y5_students,
            min_val=1000, max_val=500_000, step=5000,
            unit="students", category="school_portfolio",
            description="Target enrollment at scale. ⚠️ Changes recalculate financial model.",
        ),
        FinancialAssumption(
            key="ts_per_student_budget",
            label="Per-Student Annual Tuition / Budget ($)",
            value=round(per_student),
            min_val=5_000, max_val=100_000, step=500,
            unit="$", category="school_portfolio",
            description="Blended annual tuition per student. ⚠️ Changes recalculate financial model.",
        ),
        FinancialAssumption(
            key="ts_num_flagship_schools",
            label="Flagship Alpha Schools",
            value=2, min_val=0, max_val=10, step=1,
            unit="schools", category="school_portfolio",
            description="Premium flagship schools ($100K+ tuition) for halo effect",
        ),
        FinancialAssumption(
            key="ts_flagship_tuition",
            label="Flagship School Tuition ($)",
            value=100_000 if not is_us_state else 50_000,
            min_val=25_000, max_val=200_000, step=5_000,
            unit="$", category="school_portfolio",
            description="Annual tuition for flagship Alpha schools",
        ),
        FinancialAssumption(
            key="ts_num_schools",
            label="Total Schools at Scale",
            value=num_schools,
            min_val=1, max_val=500, step=5,
            unit="schools", category="school_portfolio",
            description="Total number of schools in the network at Year 5",
        ),

        # --- Counterparty Commitments ---
        FinancialAssumption(
            key="ts_scholarship_backstop_annual",
            label="Scholarship / Tuition Backstop ($M/year)",
            value=max(1, backstop_annual),
            min_val=0, max_val=2000, step=10,
            unit="$M", category="commitments",
            description="Annual backstop: 50% capacity × tuition (per UAE formula)",
        ),
        FinancialAssumption(
            key="ts_backstop_years",
            label="Backstop Commitment Period (years)",
            value=5, min_val=3, max_val=10, step=1,
            unit="years", category="commitments",
            description="Duration of the scholarship/tuition backstop commitment",
        ),
        FinancialAssumption(
            key="ts_launch_capital",
            label="Launch / Marketing Capital ($M)",
            value=max(5, round(y5_revenue * 0.05 / 1_000_000)),
            min_val=1, max_val=500, step=5,
            unit="$M", category="commitments",
            description="Counterparty's commitment to launch capital and marketing",
        ),
        FinancialAssumption(
            key="ts_capex_per_school",
            label="CapEx per School Buildout ($M)",
            value=round(capex_per_school / 1_000_000, 1),
            min_val=0.5, max_val=20, step=0.5,
            unit="$M", category="school_portfolio",
            description="Capital expenditure per new school. ⚠️ Changes recalculate financial model.",
        ),
    ]

    return FinancialAssumptions(assumptions=assumptions)


def get_financial_model_adjustments(
    term_sheet_adjustments: dict[str, float],
) -> dict[str, float]:
    """Given term sheet adjustments, return any that require financial model recalculation.

    Returns a dict of financial model assumption key → new value.
    """
    fm_adjustments = {}
    for ts_key, new_val in term_sheet_adjustments.items():
        if ts_key in TS_TO_FM_KEY_MAP:
            fm_key = TS_TO_FM_KEY_MAP[ts_key]
            # Convert $M fields back to raw dollars for the financial model
            if ts_key in ("ts_upfront_ip_fee",):
                new_val = new_val  # Keep as $M — the financial model stores this as $M too
            if ts_key == "ts_capex_per_school":
                new_val = new_val * 1_000_000  # Convert from $M to $
            fm_adjustments[fm_key] = new_val
    return fm_adjustments


# ---------------------------------------------------------------------------
# LLM prompts for term sheet sections
# ---------------------------------------------------------------------------

TERM_SHEET_PROMPT = """You are a senior M&A lawyer at Sullivan & Cromwell drafting a non-binding
indicative term sheet for a strategic education partnership between 2hr Learning (Alpha Holdings)
and {target}.

{uae_reference}

Context:
{context}

DEAL PARAMETERS (confirmed by user):
{deal_params}

Write a PROFESSIONAL TERM SHEET with the following exact sections. Use crisp legal/deal language.
Every term should be SPECIFIC — no placeholder ranges, commit to the exact numbers from the deal parameters.

## PARTIES
- Alpha Holdings / 2hr Learning ("Alpha")
- The counterparty (government ministry, sovereign entity, or private partner)

## TRANSACTION OVERVIEW
- 2-3 sentence description of the deal

## KEY COMMERCIAL TERMS

### Partnership Structure
- Entity type: {partnership_type}
- Alpha ownership/control: {alpha_ownership}%
- Governance (board composition, voting rights, reserved matters)

### IP Licensing & Technology
- Timeback platform license: {timeback_pct}% of per-student revenue (minimum ${min_timeback_per_student:,.0f}/student)
- AlphaCore curriculum license: included
- Incept eduLLM localisation: included
- Guide School teacher training: included
- Upfront IP development & licensing fee: ${upfront_ip:,.0f}

### Management Fee
- Alpha management fee: {mgmt_pct}% of gross school revenue (minimum ${min_mgmt_per_student:,.0f}/student)
- Payable: quarterly in arrears
- Prepayment at signing: {prepay_years} years (${prepay_amount:,.0f})

### School Portfolio
- {num_flagship} Flagship Alpha School(s) at ${flagship_tuition:,.0f}/student
- {num_schools} total schools at scale
- Year 5 target: {students_y5:,} students
- Per-student annual budget: ${per_student:,.0f}

### Financial Projections (Indicative)
{financial_summary}

### Rollout & Milestones
{rollout_info}

## COUNTERPARTY COMMITMENTS
- Scholarship / tuition backstop: ${backstop_annual:,.0f}M per year for {backstop_years} years
- Launch & marketing capital: ${launch_capital:,.0f}M
- Total backstop commitment: ${total_backstop:,.0f}M over {backstop_years} years
- School facility buildout: via development partners or counterparty (${capex_per_school:,.1f}M per school)

## ALPHA COMMITMENTS
- Three commitments: (1) Children will love school, (2) Children will learn 2x faster,
  (3) Children will develop life skills for the AI age
- Outcomes measurement: third-party verified (NWEA MAP, standardised assessments)
- Cultural IP layer: local identity, language, and values fully integrated
- Dedicated market entry team for first 24 months
- 95th percentile Guide (teacher) training via Alpha Guide School

## EXCLUSIVITY
- Alpha grants territorial exclusivity for the education model in {target}
- Duration: {exclusivity_years} years, co-terminus with partnership agreement
- Alpha retains global IP ownership

## TERM & EXIT
- Initial term: {term_years} years (renewable)
- Exit provisions: tag-along, drag-along, ROFR on partner share
- Termination for cause with 12-month cure period

## GOVERNING LAW
- [Appropriate jurisdiction for {target}]

## CONDITIONS PRECEDENT
- Regulatory approvals for school licensing
- Government endorsement / MOU
- Completion of due diligence
- Board approvals from both parties

## CONFIDENTIALITY
- This term sheet is confidential and non-binding (other than confidentiality and exclusivity)
- Subject to definitive agreement execution

---

*This Indicative Term Sheet is for discussion purposes only and does not constitute a binding offer or commitment. All terms are subject to final negotiation and execution of definitive agreements.*
"""


async def generate_term_sheet(
    target: str,
    country_profile: CountryProfile,
    education_analysis: EducationAnalysis,
    strategy: Strategy,
    financial_model: FinancialModel,
    assumptions: FinancialAssumptions,
    term_sheet_assumptions: FinancialAssumptions | None = None,
) -> tuple[str, str]:
    """Generate a term sheet.

    Returns (term_sheet_markdown, term_sheet_docx_path).
    """
    logger.info("Generating term sheet for %s", target)

    # Get term sheet deal parameters from assumptions (if available)
    ts = {}
    if term_sheet_assumptions and term_sheet_assumptions.assumptions:
        ts = {item.key: item.value for item in term_sheet_assumptions.assumptions}

    # Build context
    context_parts = []
    if country_profile.demographics.total_population:
        context_parts.append(f"Population: {country_profile.demographics.total_population:,.0f}")
    if country_profile.education.k12_enrolled:
        context_parts.append(f"K-12 students: {country_profile.education.k12_enrolled:,.0f}")
    if country_profile.economy.gdp_per_capita:
        context_parts.append(f"GDP/capita: ${country_profile.economy.gdp_per_capita:,.0f}")
    if country_profile.target.tier:
        context_parts.append(f"Tier: {country_profile.target.tier}")
    if strategy.entry_mode:
        context_parts.append(f"Entry mode: {strategy.entry_mode.value}")
    if strategy.partnership_structure.type:
        context_parts.append(f"Partnership: {strategy.partnership_structure.type.value}")

    # Deal parameters from term sheet assumptions
    partnership_map = {1: "Joint Venture (JV)", 2: "Licensing Agreement", 3: "Direct Operation"}
    partnership_type = partnership_map.get(int(ts.get("ts_partnership_type", 1)), "Joint Venture (JV)")
    alpha_ownership = int(ts.get("ts_alpha_ownership_pct", 49))
    term_years = int(ts.get("ts_term_years", 25))
    exclusivity_years = int(ts.get("ts_exclusivity_years", 25))
    upfront_ip = ts.get("ts_upfront_ip_fee", round(financial_model.upfront_ip_fee / 1_000_000)) * 1_000_000
    mgmt_pct = int(ts.get("ts_management_fee_pct", round(financial_model.management_fee_pct * 100)))
    timeback_pct = int(ts.get("ts_timeback_license_pct", round(financial_model.timeback_license_pct * 100)))
    prepay_years = int(ts.get("ts_mgmt_fee_prepay_years", 2))
    prepay_amount = ts.get("ts_mgmt_fee_prepay_amount", 10) * 1_000_000
    students_y5 = int(ts.get("ts_students_year5", financial_model.pnl_projection[-1].students if financial_model.pnl_projection else 100_000))
    per_student = ts.get("ts_per_student_budget", 25_000)
    num_flagship = int(ts.get("ts_num_flagship_schools", 2))
    flagship_tuition = ts.get("ts_flagship_tuition", 100_000)
    num_schools = int(ts.get("ts_num_schools", 10))
    backstop_annual = ts.get("ts_scholarship_backstop_annual", 100)
    backstop_years = int(ts.get("ts_backstop_years", 5))
    launch_capital = ts.get("ts_launch_capital", 50)
    capex_per_school = ts.get("ts_capex_per_school", 5)

    # Build deal params string for context
    deal_params = f"""- Partnership: {partnership_type}
- Alpha ownership: {alpha_ownership}%
- Upfront IP fee: ${upfront_ip:,.0f}
- Management fee: {mgmt_pct}% of revenue
- Timeback license: {timeback_pct}% of per-student budget
- Prepayment: {prepay_years} years (${prepay_amount:,.0f})
- Year 5 students: {students_y5:,}
- Per-student budget: ${per_student:,.0f}
- Flagship schools: {num_flagship} at ${flagship_tuition:,.0f}
- Total schools at scale: {num_schools}
- Backstop: ${backstop_annual:,.0f}M/yr for {backstop_years} years
- Launch capital: ${launch_capital:,.0f}M
- Term: {term_years} years"""

    # Financial info from model
    fin_lines = []
    if financial_model.pnl_projection:
        for p in financial_model.pnl_projection:
            fin_lines.append(
                f"- Year {p.year}: {p.students:,} students | {p.schools} schools | "
                f"${p.revenue:,.0f} revenue | ${p.ebitda:,.0f} EBITDA"
            )
    financial_summary = "\n".join(fin_lines) or "- See attached financial model"

    rollout_lines = []
    if strategy.phased_rollout:
        for ph in strategy.phased_rollout[:5]:
            rollout_lines.append(
                f"- {ph.phase} ({ph.timeline}): {ph.student_count:,} students"
                if ph.student_count else f"- {ph.phase} ({ph.timeline})"
            )
    rollout_info = "\n".join(rollout_lines) or "- Phased rollout over 5 years"

    # Min per-student fees (based on UAE: $2,500 mgmt, $5,000 timeback)
    min_mgmt_per_student = max(2_500, round(per_student * mgmt_pct / 100))
    min_timeback_per_student = max(5_000, round(per_student * timeback_pct / 100))

    term_sheet_md = await call_llm_plain(
        system_prompt=TERM_SHEET_PROMPT.format(
            target=target,
            target_upper=target.upper(),
            uae_reference=UAE_HOT_REFERENCE,
            context="\n".join(context_parts),
            deal_params=deal_params,
            partnership_type=partnership_type,
            alpha_ownership=alpha_ownership,
            timeback_pct=timeback_pct,
            mgmt_pct=mgmt_pct,
            upfront_ip=upfront_ip,
            min_mgmt_per_student=min_mgmt_per_student,
            min_timeback_per_student=min_timeback_per_student,
            prepay_years=prepay_years,
            prepay_amount=prepay_amount,
            num_flagship=num_flagship,
            flagship_tuition=flagship_tuition,
            num_schools=num_schools,
            students_y5=students_y5,
            per_student=per_student,
            financial_summary=financial_summary,
            rollout_info=rollout_info,
            backstop_annual=backstop_annual,
            backstop_years=backstop_years,
            total_backstop=backstop_annual * backstop_years,
            launch_capital=launch_capital,
            capex_per_school=capex_per_school,
            exclusivity_years=exclusivity_years,
            term_years=term_years,
        ),
        user_prompt=f"Produce the complete term sheet for the 2hr Learning × {target} partnership.",
    )

    # Save as DOCX
    docx_path = _build_term_sheet_docx(target, term_sheet_md)

    logger.info("Term sheet generated for %s", target)
    return term_sheet_md, docx_path


def _build_term_sheet_docx(target: str, markdown: str) -> str:
    """Render the term sheet markdown into a professional DOCX."""
    doc = DocxDocument()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = DocxPt(11)
    style.font.color.rgb = DocxRGB(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = DocxPt(4)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, (size, color, bold) in {
        1: (DocxPt(20), DocxRGB(0x1a, 0x1a, 0x2e), True),
        2: (DocxPt(14), DocxRGB(0x00, 0x6D, 0x77), True),
        3: (DocxPt(12), DocxRGB(0x33, 0x33, 0x33), True),
    }.items():
        hs = doc.styles[f"Heading {level}"]
        hs.font.size = size
        hs.font.color.rgb = color
        hs.font.bold = bold
        hs.font.name = "Calibri"
        hs.paragraph_format.space_before = DocxPt(14)
        hs.paragraph_format.space_after = DocxPt(6)

    # --- Header ---
    doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("INDICATIVE TERM SHEET")
    run.font.size = DocxPt(24)
    run.font.color.rgb = DocxRGB(0x1a, 0x1a, 0x2e)
    run.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(f"2hr Learning (Alpha) × {target}")
    run.font.size = DocxPt(16)
    run.font.color.rgb = DocxRGB(0x00, 0x6D, 0x77)
    run.bold = True

    subtitle_p2 = doc.add_paragraph()
    subtitle_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p2.add_run("Strategic Education Partnership")
    run.font.size = DocxPt(12)
    run.font.color.rgb = DocxRGB(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = conf_p.add_run("CONFIDENTIAL & NON-BINDING")
    run.font.size = DocxPt(11)
    run.font.color.rgb = DocxRGB(0xCC, 0x00, 0x00)
    run.bold = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%B %Y"))
    run.font.size = DocxPt(10)
    run.font.color.rgb = DocxRGB(0x99, 0x99, 0x99)

    doc.add_page_break()

    # --- Render markdown body ---
    for line in markdown.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            text = stripped[2:]
            if text.startswith("**") and "**" in text[2:]:
                end = text.index("**", 2)
                bold_part = text[2:end]
                rest = text[end + 2:]
                p = doc.add_paragraph(style="List Bullet")
                r1 = p.add_run(bold_part)
                r1.bold = True
                p.add_run(rest)
            else:
                doc.add_paragraph(text, style="List Bullet")
        elif stripped.startswith("*") and stripped.endswith("*"):
            # Italic disclaimer
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = DocxPt(9)
            run.font.color.rgb = DocxRGB(0x99, 0x99, 0x99)
        elif stripped.startswith("|"):
            # Simple table row — render as plain text for term sheets
            doc.add_paragraph(stripped)
        elif stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = DocxPt(8)
            p.paragraph_format.space_after = DocxPt(8)
        else:
            p = doc.add_paragraph()
            # Handle inline bold
            _add_formatted_text(p, stripped)

    # Save
    output_dir = os.path.join(OUTPUT_DIR, target.lower().replace(" ", "_"))
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"{target.replace(' ', '_')}_term_sheet.docx")
    doc.save(path)
    logger.info("Term sheet DOCX saved: %s", path)
    return path


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text with **bold** handling."""
    pos = 0
    while pos < len(text):
        bold_start = text.find("**", pos)
        if bold_start != -1:
            bold_end = text.find("**", bold_start + 2)
            if bold_end != -1:
                if bold_start > pos:
                    paragraph.add_run(text[pos:bold_start])
                run = paragraph.add_run(text[bold_start + 2:bold_end])
                run.bold = True
                pos = bold_end + 2
                continue
        paragraph.add_run(text[pos:])
        break
