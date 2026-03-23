"""Strategy Agent — country-specific pitch narrative.

Two LLM calls:
  1. Structured Strategy extraction (schema-driven)
  2. Combined strategy brief (~3-4 pages)

The deal model is FIXED. This agent's job is to conform the pitch to
the country — change the words, not the numbers.

All outputs suitable for direct presentation to heads of state.
"""

from __future__ import annotations

import logging
import os

from docx import Document as DocxDocument
from docx.shared import Pt as DocxPt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import (
    Strategy, CountryProfile, EducationAnalysis,
    EntryMode, TargetType,
)
from services.llm import call_llm, call_llm_plain
from config import OUTPUT_DIR

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Head-of-state preamble — prepended to every LLM prompt
# ---------------------------------------------------------------------------

HEAD_OF_STATE_PREAMBLE = """CRITICAL — EXTERNALLY PRESENTABLE ARTIFACTS:
These documents will be presented DIRECTLY to heads of state, sovereign rulers, royal family
members, and senior government ministers. The UAE version of these materials was personally
reviewed by the daughter of the #2 figure in Abu Dhabi and the #3 figure in the UAE.

ABSOLUTE RULES:
1. NEVER include internal classifications (tiers, scores, ratings, internal frameworks)
2. NEVER reference "Tier 1", "Tier 2", "Tier 3" or any scoring/ranking system
3. NEVER include analytical frameworks, scoring matrices, or internal assessment tools
4. NEVER reference PPP adjustments, GDP scaling, or pricing methodology
5. NEVER use phrases like "Alpha Relevance", "fit score", or internal jargon
6. Write in the voice of a trusted advisor presenting to royalty — measured, confident, specific
7. All financial figures are FIXED (the deal model does not change by country)
8. Research is for cultural context and narrative color only — it does NOT drive financial numbers
9. Use the country's formal diplomatic name and proper honorifics for leaders
10. Every artifact must be suitable for immediate external presentation without editing
"""

# ---------------------------------------------------------------------------
# Single strategy report prompt (replaces 3 section prompts)
# ---------------------------------------------------------------------------

STRATEGY_REPORT_PROMPT = HEAD_OF_STATE_PREAMBLE + """
You are a trusted senior advisor preparing a strategy brief for **{target}**
to support an education partnership with Alpha Holdings, Inc.

Conform this pitch to the country. The model is fixed. Your job is to change
the words to feel right for the country.

The deal structure is FIXED:
- Operator & Licensor model (Marriott model): counterparty owns 100%, Alpha owns 0%
- $750M fixed development costs ($250M AlphaCore + $250M App R&D + $250M LifeSkills)
- $25K fixed national per-student budget, 100K student-year minimum
- Flagship tuition $40K-$100K (set by AGI of top 20% families)
- Management fee 10%, Timeback license 20% (both non-negotiable)
DO NOT derive or recommend deal structures — they are set.

Research is color commentary. Use it to make the pitch feel country-specific,
not to derive financial numbers.

## Country & Education Context
{country_context}

{education_context}

## Write a 3-4 page strategy brief with these sections:

# Strategy Brief: {target}

## Partnership Narrative
Why this country, why now. Alignment with national vision and leadership priorities.
What makes this moment right for a transformative education partnership.

## Dual-School Positioning
How flagship schools (2-3 premium campuses, $40K-$100K tuition) and national schools
($25K budget, 100K+ students) fit this specific market. How flagship success in the
capital proves the model and unlocks national scale.

## Cultural Localisation
Life-skills ideas drawn from this country's heritage and values. Cultural IP layer
concept (like AsasOne in the UAE). Language of instruction considerations.
Curriculum adaptation needs.

## Regulatory Pathway
Key steps to operate as a foreign education partner. Key decision-makers and
government bodies. Realistic timeline estimate from agreement to first school.

## Competitive Positioning
Brief landscape of existing private/international education. How Alpha's integrated
model (Timeback + AlphaCore + Guide School + eduLLM) differentiates. Keep this
concise — no exhaustive competitor tables.

RULES:
- Do NOT include TAM/SAM/SOM calculations
- Do NOT include IRR/MOIC projections or 5-year financial tables
- Do NOT include entry mode comparison matrices or scoring frameworks
- Do NOT include risk registers or probability/impact matrices
- Do NOT include PPP-adjusted pricing
- Keep the tone measured, confident, and suitable for a head-of-state audience
- Cite specific country data where it adds color
"""

US_STATE_STRATEGY_PROMPT = HEAD_OF_STATE_PREAMBLE + """
You are a senior education policy strategist preparing a strategy brief
for **{target}** to support Alpha Holdings, Inc.'s US expansion.

## Context
{country_context}

{education_context}

## Write a concise strategy brief:

# Strategy Brief: {target}

## Opportunity Narrative
Why this state, why now. Political climate for school choice. Key pain points
in the current system that Alpha addresses.

## Market Positioning
How Alpha's model fits the state's education landscape, ESA/voucher programs,
and parent demand for alternatives.

## Regulatory & Political Pathway
Key steps, decision-makers, and timeline for operating in this state.

## Competitive Landscape
Brief overview of existing alternatives and Alpha's differentiation.

RULES:
- Do NOT compute financial model numbers
- Keep the tone professional and policy-focused
- Cite specific state data where relevant
"""

REPORT_REVISION_PROMPT = HEAD_OF_STATE_PREAMBLE + """You are revising a strategy brief based on feedback.

Original strategy:
{original_report}

Feedback:
{feedback}

Produce a revised version of the FULL brief incorporating the feedback.
Maintain the same structure, tone, and conciseness."""


async def run_strategy(
    target: str,
    country_profile: CountryProfile,
    education_analysis: EducationAnalysis,
    entry_mode: EntryMode | None = None,
    feedback: str | None = None,
    previous_report: str | None = None,
) -> tuple[Strategy, str, str]:
    """Execute strategy development. Returns (strategy, report_md, docx_path)."""
    logger.info("Running strategy agent for %s", target)

    country_ctx = _build_country_context(country_profile)
    education_ctx = _build_education_context(education_analysis)

    # --- Detect US state ---
    is_us_state = country_profile.target.type == TargetType.US_STATE

    # --- Minimum student target ---
    # Sovereign nations: 100K student-year minimum for national schools
    # US states: 10% of school-age population
    if is_us_state:
        school_age_pop = (
            country_profile.demographics.population_0_18
            or country_profile.education.k12_enrolled
            or 500_000
        )
        min_y5_students = max(5_000, int(school_age_pop * 0.10))
    else:
        min_y5_students = 100_000  # Workshop decision: 100K student-year minimum

    # --- Structured strategy ---
    if is_us_state:
        structured_prompt = (
            HEAD_OF_STATE_PREAMBLE +
            "Extract a structured Strategy for a US state market-entry deal.\n\n"
            "FIXED DEAL PARAMETERS:\n"
            "- Alpha is exclusive operator & licensor; local entity owns 100% equity\n"
            f"- Year 5 student target: minimum {min_y5_students:,}\n"
            "- Management fee: 10%, Timeback license: 20% (non-negotiable)\n\n"
            "SCHOOL TYPE NAMING:\n"
            "- Flagship schools may be called 'Alpha Flagship School' (these are 100% Alpha-owned)\n"
            "- State-owned/public schools MUST NOT have 'Alpha' in their name. "
            "Use a name like 'National Program', 'State School', or similar. "
            "It is OK to say they are 'operated by Alpha'.\n\n"
            "Set partnership_structure to operator_licensor.\n"
            "Set ownership_split to '0/100 — Alpha operates; local entity owns 100%'.\n"
            "Produce a structured strategy."
        )
    else:
        structured_prompt = (
            HEAD_OF_STATE_PREAMBLE +
            "Extract a structured Strategy for a sovereign education partnership.\n\n"
            "FIXED DEAL PARAMETERS:\n"
            "- Flagship: $40K-$100K tuition, 2-3 schools, capital cities\n"
            f"- National: $25K fixed budget, {min_y5_students:,} student-year minimum\n"
            "- Equity: 100/0 — Counterparty 100%, Alpha 0% (Operator & Licensor)\n"
            "- Management fee: 10%, Timeback: 20% (non-negotiable)\n"
            "- Fixed development: $750M ($250M × 3, non-negotiable)\n\n"
            "SCHOOL TYPE NAMING:\n"
            "- Flagship schools should be called 'Alpha Flagship School' (these are 100% Alpha-owned)\n"
            "- National/country-owned schools MUST NOT have 'Alpha' in their name. "
            "Use a culturally resonant name like 'National Program', 'National School', or similar. "
            "It is OK to say they are 'operated by Alpha'.\n\n"
            "Set partnership_structure to operator_licensor.\n"
            "Set ownership_split to '100/0 — Counterparty owns 100%, Alpha is operator & licensor'.\n"
            "Produce a structured strategy."
        )
    try:
        strategy: Strategy = await call_llm(
            system_prompt=structured_prompt,
            user_prompt=(
                f"Target: {target}\nCountry: {country_ctx}\nEducation: {education_ctx}\n"
                f"Entry mode preference: {entry_mode.value if entry_mode else 'not specified'}"
            ),
            output_schema=Strategy,
        )
    except Exception as exc:
        logger.warning("Structured strategy failed: %s", exc)
        strategy = Strategy(entry_mode=entry_mode or EntryMode.HYBRID)

    if entry_mode:
        strategy.entry_mode = entry_mode

    # --- Enforce non-negotiable deal parameters on the structured object ---
    if not is_us_state:
        strategy.partnership_structure.ownership_split = (
            "100/0 — Counterparty owns 100%, Alpha is exclusive operator & licensor"
        )
    else:
        strategy.partnership_structure.ownership_split = (
            "0/100 — Alpha operates as exclusive operator & licensor; local entity owns 100%"
        )

    # Enforce Y5 student floor
    if strategy.target_student_count_year5 and strategy.target_student_count_year5 < min_y5_students:
        logger.info(
            "Overriding Y5 student target from %s to %s (minimum floor)",
            strategy.target_student_count_year5, min_y5_students,
        )
        strategy.target_student_count_year5 = min_y5_students
    elif not strategy.target_student_count_year5:
        strategy.target_student_count_year5 = min_y5_students

    # --- Narrative report ---
    if feedback and previous_report:
        report_md = await call_llm_plain(
            system_prompt=REPORT_REVISION_PROMPT.format(
                original_report=previous_report, feedback=feedback
            ),
            user_prompt=f"Revise the strategy brief for {target}.",
        )
    else:
        if is_us_state:
            prompt = US_STATE_STRATEGY_PROMPT.format(
                target=target, country_context=country_ctx, education_context=education_ctx,
            )
        else:
            prompt = STRATEGY_REPORT_PROMPT.format(
                target=target, country_context=country_ctx, education_context=education_ctx,
            )

        logger.info("Generating strategy brief for %s", target)
        report_md = await call_llm_plain(
            system_prompt=prompt,
            user_prompt=f"Write the strategy brief for {target}.",
        )

    # --- Save DOCX ---
    docx_path = _save_report_docx(target, report_md, "Market Entry Strategy")

    logger.info("Strategy complete for %s (mode=%s)", target, strategy.entry_mode)
    return strategy, report_md, docx_path


def _build_country_context(p: CountryProfile) -> str:
    parts = [f"**Country Profile: {p.target.name}**"]
    if p.demographics.total_population:
        parts.append(f"- Population: {p.demographics.total_population:,.0f}")
    if p.demographics.population_0_18:
        parts.append(f"- School-age population: {p.demographics.population_0_18:,.0f}")
    if p.economy.gdp:
        parts.append(f"- GDP: ${p.economy.gdp:,.0f}")
    if p.economy.gdp_per_capita:
        parts.append(f"- GDP per capita: ${p.economy.gdp_per_capita:,.0f}")
    if p.economy.gdp_growth_rate:
        parts.append(f"- GDP growth: {p.economy.gdp_growth_rate}%")
    if p.economy.sovereign_wealth_fund:
        parts.append(f"- SWF: {p.economy.sovereign_wealth_fund}")
    if p.education.k12_enrolled:
        parts.append(f"- K-12 students: {p.education.k12_enrolled:,.0f}")
    if p.education.avg_private_tuition:
        parts.append(f"- Avg private tuition: ${p.education.avg_private_tuition:,.0f}")
    if p.regulatory.foreign_ownership_rules:
        parts.append(f"- Foreign ownership: {p.regulatory.foreign_ownership_rules}")
    if p.political_context.national_vision_plan:
        parts.append(f"- National vision: {p.political_context.national_vision_plan}")
    if p.competitive_landscape.major_operators:
        ops = ", ".join(o.name for o in p.competitive_landscape.major_operators[:5])
        parts.append(f"- Major operators: {ops}")
    return "\n".join(parts)


def _build_education_context(e: EducationAnalysis) -> str:
    parts = ["**Education Analysis Summary:**"]
    if e.system_diagnosis.primary_pain_points:
        parts.append(f"- Student pain points: {'; '.join(e.system_diagnosis.primary_pain_points[:5])}")
    if e.system_diagnosis.parent_pain_points:
        parts.append(f"- Parent pain points: {'; '.join(e.system_diagnosis.parent_pain_points[:5])}")
    if e.system_diagnosis.government_pain_points:
        parts.append(f"- Government pain points: {'; '.join(e.system_diagnosis.government_pain_points[:5])}")
    if e.reform_landscape.active_reforms:
        parts.append(f"- Active reforms: {'; '.join(e.reform_landscape.active_reforms[:5])}")
    if e.two_hr_learning_fit.unique_value_propositions:
        parts.append(f"- Alpha UVPs: {'; '.join(e.two_hr_learning_fit.unique_value_propositions[:5])}")
    if e.two_hr_learning_fit.model_recommendation:
        parts.append(f"- Recommended model: {e.two_hr_learning_fit.model_recommendation.value}")
    return "\n".join(parts) or "Education analysis pending"


def _save_report_docx(target: str, markdown: str, title: str) -> str:
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = DocxPt(11)
    h = doc.add_heading(f"{title}: {target}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("CONFIDENTIAL & PROPRIETARY — Alpha Holdings, Inc.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    for line in markdown.split("\n"):
        s = line.strip()
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s.startswith("|"):
            doc.add_paragraph(s)
        elif s:
            doc.add_paragraph(s)
    output_dir = os.path.join(OUTPUT_DIR, target.lower().replace(" ", "_"))
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"{target.replace(' ', '_')}_strategy_report.docx")
    doc.save(path)
    return path
