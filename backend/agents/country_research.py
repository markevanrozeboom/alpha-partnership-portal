"""Country & Education Research Agent — combined country + education profile.

Produces a single research report covering macroeconomic context, education
system analysis, and cultural/regulatory landscape.  Two LLM calls:
  1. Structured data extraction → CountryProfile + EducationAnalysis
  2. Combined narrative report   → ~500-word focused briefing

Education research (previously a separate agent) is merged here so the
pipeline makes ONE Perplexity call for general research and ONE for education,
then synthesises everything in a single report.

All outputs are suitable for direct presentation to heads of state.
"""

from __future__ import annotations

import asyncio
import logging
import os

from docx import Document as DocxDocument
from docx.shared import Pt as DocxPt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import (
    CountryProfile, TargetInfo, TargetType,
    Demographics, Economy, EducationData, Regulatory, PoliticalContext,
    FlagshipMarketData,
)
from services.llm import call_llm, call_llm_plain
from services.perplexity import (
    research_country, research_education, research_us_state,
)
from services.world_bank import get_country_data
from config import OUTPUT_DIR
from config.rules_loader import (
    get_state_spending_data,
    get_spending_spotlight_national_trends,
    get_spending_spotlight_alpha_insights,
)

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Head-of-state preamble — prepended to every LLM prompt
# ---------------------------------------------------------------------------

HEAD_OF_STATE_PREAMBLE = """CRITICAL — EXTERNALLY PRESENTABLE ARTIFACTS:
These documents will be presented DIRECTLY to heads of state, sovereign rulers, royal family
members, and senior government ministers. The UAE version of these materials was personally
reviewed by the daughter of the #2 figure in Abu Dhabi and the #3 figure in the UAE.

ABSOLUTE RULES:
1. NEVER include internal classifications (tiers, scores, ratings, internal frameworks)
2. NEVER reference "Tier 1", "Tier 2", "Tier 3" or any scoring/ranking system
3. NEVER include analytical frameworks, scoring matrices, or internal assessment tools
4. NEVER reference PPP adjustments, GDP scaling, or pricing methodology
5. NEVER use phrases like "Alpha Relevance", "fit score", or internal jargon
6. Write in the voice of a trusted advisor presenting to royalty — measured, confident, specific
7. All financial figures are FIXED (the deal model does not change by country)
8. Research is for cultural context and narrative color only — it does NOT drive financial numbers
9. Use the country's formal diplomatic name and proper honorifics for leaders
10. Every artifact must be suitable for immediate external presentation without editing
"""

# ---------------------------------------------------------------------------
# US state list (for target-type detection)
# ---------------------------------------------------------------------------

US_STATES = {
    "alabama", "alaska", "arizona", "arkansas", "california", "colorado",
    "connecticut", "delaware", "florida", "georgia", "hawaii", "idaho",
    "illinois", "indiana", "iowa", "kansas", "kentucky", "louisiana",
    "maine", "maryland", "massachusetts", "michigan", "minnesota",
    "mississippi", "missouri", "montana", "nebraska", "nevada",
    "new hampshire", "new jersey", "new mexico", "new york",
    "north carolina", "north dakota", "ohio", "oklahoma", "oregon",
    "pennsylvania", "rhode island", "south carolina", "south dakota",
    "tennessee", "texas", "utah", "vermont", "virginia", "washington",
    "west virginia", "wisconsin", "wyoming", "district of columbia",
}


def _detect_target_type(target: str) -> TargetType:
    if target.lower().strip() in US_STATES:
        return TargetType.US_STATE
    return TargetType.SOVEREIGN_NATION


# ---------------------------------------------------------------------------
# LLM call 1 — structured data extraction
# ---------------------------------------------------------------------------

SYNTHESIS_PROMPT = HEAD_OF_STATE_PREAMBLE + """
You are a senior research analyst synthesising country and education data
into a structured profile. Extract specific numbers wherever available.
If a data point is not available, provide your best estimate and note it.

KEY DATA POINTS TO EXTRACT:
1. School-age population (0-18)
2. Government per-student funding (public spend per pupil)
3. Metropolitan wealth indicators / AGI top-20% household income
4. Regulatory basics (foreign ownership rules, licensing process, timeline)
5. Language of instruction

Also extract any education system pain points, reform initiatives, and government appetite
for foreign education models — these will be used to seed the education analysis stage.

Output a JSON matching the CountryProfile schema."""

# ---------------------------------------------------------------------------
# LLM call 1b — flagship market data extraction
# ---------------------------------------------------------------------------

FLAGSHIP_EXTRACTION_PROMPT = """You are a senior research analyst extracting \
structured metro-level data for premium school sizing.

Given the research text below, extract data for each of the TOP 3 METROS \
(by population) in the country.

For EACH metro provide:
- metro_name: name of the metropolitan area
- is_capital: true if this metro contains the national capital
- metro_population: estimated metro area population (integer)
- k12_children: estimated K-12 age children (ages 5-18) in the metro (integer)
- children_in_families_income_above_200k: estimated K-12 children living in \
families with annual household income ≥ $200,000 USD (integer). Use PPP \
conversion if the research provides local-currency data.
- children_in_families_income_above_500k: estimated K-12 children living in \
families with annual household income ≥ $500,000 USD (integer). Use PPP \
conversion if the research provides local-currency data.
- most_expensive_nonboarding_tuition: annual tuition in USD of the most \
expensive NON-BOARDING private school in this metro (float). If only local \
currency is given, convert to USD.
- most_expensive_nonboarding_school: name of that school (string)

Also provide:
- country_most_expensive_nonboarding_tuition: the highest non-boarding \
private school annual tuition in the entire country (float, USD)
- country_most_expensive_nonboarding_school: name of that school (string)

IMPORTANT:
- Provide your best numerical ESTIMATE for every field. Never leave zeros \
unless you truly believe the answer is zero.
- For income thresholds, reason from the country's income distribution, \
Gini coefficient, and wealth data.
- Order metros by population (largest first).
- Output valid JSON matching the FlagshipMarketData schema.
"""

# ---------------------------------------------------------------------------
# LLM call 2 — combined narrative report (~500 words)
# ---------------------------------------------------------------------------

REPORT_PROMPT = HEAD_OF_STATE_PREAMBLE + """
You are a trusted senior advisor preparing a confidential country briefing
for **{target}** to support a potential education partnership with
2hr Learning (Alpha).

## Context Data
{data_context}

## Instructions

Write a concise, authoritative country and education briefing (~500 words).
Structure it as follows:

# Country & Education Briefing: {target}

## National Context
2-3 paragraphs covering the macroeconomic position, demographic profile,
and political/regulatory environment relevant to a large-scale education
partnership. Reference the country's formal name and leadership appropriately.

## Education Landscape
2-3 paragraphs covering:
- School-age population and enrollment figures
- Government per-student funding levels
- Current education system performance and key gaps
- Language of instruction and curriculum framework
- Reform appetite and any active government education initiatives

## Partnership Environment
1-2 paragraphs covering:
- Regulatory pathway for foreign education operators
- Metropolitan wealth indicators and capacity for premium education
- Cultural context relevant to education innovation
- Precedent for international education partnerships

## Key Considerations
A short closing paragraph highlighting the 3-5 most important factors
for Alpha's leadership team to consider.

RULES:
- Use research data for narrative context and color ONLY
- Do NOT compute financial model numbers (those are fixed and handled elsewhere)
- Do NOT include TAM/SAM/SOM calculations
- Do NOT include risk registers or scoring matrices
- Do NOT reference tiers, PPP adjustments, or internal frameworks
- Keep the tone measured and confident — suitable for a head-of-state audience
- Cite specific numbers with sources where available
- Maximum ~500 words total
"""

US_STATE_REPORT_PROMPT = HEAD_OF_STATE_PREAMBLE + """
You are a senior education policy advisor preparing a state briefing
for **{target}** to support 2hr Learning (Alpha)'s US expansion strategy.

## Context Data
{data_context}

## Instructions

Write a concise state education briefing (~500 words).

# Education Briefing: {target}

## State Overview
1-2 paragraphs on the state's economic position, demographics, and
political orientation toward education choice.

## Education System
2-3 paragraphs covering:
- K-12 enrollment and spending per pupil (vs. national average)
- NAEP scores and outcome trends
- Teacher workforce (salary, student-teacher ratio)
- Key performance gaps and areas of concern

## School Choice & ESA Landscape
1-2 paragraphs on:
- ESA/voucher programs: program name, amount, eligibility, enrollment
- Charter school landscape
- Political appetite for education innovation and parental choice

## Alpha Opportunity
1 paragraph summarising the key factors that make this state
relevant for Alpha's US model.

RULES:
- Use Spending Spotlight data for factual context
- Do NOT compute financial model numbers
- Keep the tone professional and policy-focused
- Cite specific numbers with sources
- Maximum ~500 words total
"""

REPORT_REVISION_PROMPT = HEAD_OF_STATE_PREAMBLE + \
    """You are revising a country/education briefing based on user feedback.

Here is the original report:

{original_report}

Here is the user's feedback:

{feedback}

Produce a revised version of the FULL report incorporating the feedback.
Maintain the same structure, tone, and conciseness. Do not expand beyond
~500 words unless the feedback specifically requests more detail."""


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

async def run_country_research(
    target: str,
    feedback: str | None = None,
    previous_report: str | None = None,
) -> tuple[CountryProfile, str, str]:
    """Execute the country research pipeline.

    Returns (profile, report_markdown, report_docx_path).

    Two LLM calls (fresh report):
      1. Structured data extraction → CountryProfile (+ EducationAnalysis stored on profile)
      2. Combined narrative report
    One LLM call for revision mode.
    """
    target_type = _detect_target_type(target)
    logger.info("Researching %s (type: %s)", target, target_type.value)

    # ------------------------------------------------------------------
    # Data gathering — Perplexity research + World Bank + Spending Spotlight
    # ------------------------------------------------------------------

    if target_type == TargetType.US_STATE:
        perplexity_result = await research_us_state(target)
        wb_data: dict = {}
        edu_research_text = ""
        edu_citations: list = []
    else:
        # Two deep-research Perplexity calls + World Bank run concurrently
        # (flagship data is extracted from these same results — no separate call)
        (
            perplexity_result, wb_data, edu_perplexity,
        ) = await asyncio.gather(
            research_country(target),
            get_country_data(target),
            research_education(target),
        )
        edu_research_text = edu_perplexity.get("answer", "")
        edu_citations = edu_perplexity.get("citations", [])

    research_text = perplexity_result.get("answer", "")
    citations = perplexity_result.get("citations", [])

    # ------------------------------------------------------------------
    # Build structured profile from World Bank data
    # ------------------------------------------------------------------
    profile = CountryProfile()
    profile.target = TargetInfo(name=target, type=target_type, region="")
    if isinstance(citations, list):
        profile.research_sources = [str(c) for c in citations]
    if isinstance(edu_citations, list):
        profile.research_sources.extend(str(c) for c in edu_citations)

    if wb_data:
        profile.demographics.total_population = wb_data.get("population")
        profile.demographics.population_0_18 = wb_data.get("population_0_14")
        profile.demographics.growth_rate = wb_data.get("population_growth")
        profile.demographics.urbanisation = wb_data.get("urbanization")
        profile.demographics.gini_coefficient = wb_data.get("gini")
        profile.economy.gdp = wb_data.get("gdp")
        profile.economy.gdp_per_capita = wb_data.get("gdp_per_capita")
        profile.economy.gdp_growth_rate = wb_data.get("gdp_growth")
        profile.economy.inflation = wb_data.get("inflation")
        profile.education.literacy_rate = wb_data.get("literacy")
        profile.education.education_budget_pct_gdp = wb_data.get("education_spend_pct_gdp")
        profile.education.student_teacher_ratio = wb_data.get("pupil_teacher_ratio_primary")

    # ------------------------------------------------------------------
    # LLM call 1: structured data extraction → CountryProfile + EducationAnalysis
    # ------------------------------------------------------------------
    try:
        synth: CountryProfile = await call_llm(
            system_prompt=SYNTHESIS_PROMPT,
            user_prompt=(
                f"Target: {target} ({target_type.value})\n"
                f"World Bank: {wb_data}\n"
                f"Research: {research_text}\n"
                f"Education Research: {edu_research_text}"
            ),
            output_schema=CountryProfile,
        )
        _merge_profiles(profile, synth)
    except Exception as exc:
        logger.warning("Structured synthesis failed: %s", exc)

    if not profile.target.region:
        profile.target.region = "North America" if target_type == TargetType.US_STATE else ""

    # ------------------------------------------------------------------
    # LLM call 1b: flagship market data extraction (sovereign only)
    #
    # Uses the country + education research already gathered above —
    # no separate Perplexity call needed. The deep-research results
    # contain city populations, income distributions, private school
    # tuitions, and UHNW data that the LLM can synthesize.
    # ------------------------------------------------------------------
    if target_type != TargetType.US_STATE:
        if not research_text and not edu_research_text:
            logger.warning(
                "No research text available for %s "
                "— flagship optimization will use fallback defaults.",
                target,
            )
        else:
            logger.info(
                "Extracting flagship market data for %s from "
                "country research (%d chars) + education research (%d chars)",
                target, len(research_text), len(edu_research_text),
            )
            try:
                fmd: FlagshipMarketData = await call_llm(
                    system_prompt=FLAGSHIP_EXTRACTION_PROMPT,
                    user_prompt=(
                        f"Target country: {target}\n\n"
                        f"Country Research (demographics, economy, "
                        f"income distribution, cities, UHNW data):\n"
                        f"{research_text[:6000]}\n\n"
                        f"Education Research (school tuitions, private "
                        f"school market, premium schools):\n"
                        f"{edu_research_text[:4000]}"
                    ),
                    output_schema=FlagshipMarketData,
                )
                profile.flagship_market_data = fmd
                logger.info(
                    "Flagship market data extracted: %d metros, "
                    "top school $%s",
                    len(fmd.metros),
                    f"{fmd.country_most_expensive_nonboarding_tuition:,.0f}",
                )
            except Exception as exc:
                logger.warning(
                    "Flagship market data extraction failed: %s", exc,
                )

    # ------------------------------------------------------------------
    # Populate US state education data from Spending Spotlight
    # ------------------------------------------------------------------
    spending_spotlight_context = ""
    if target_type == TargetType.US_STATE:
        ss_data = get_state_spending_data(target)
        national = get_spending_spotlight_national_trends()
        alpha_insights = get_spending_spotlight_alpha_insights()
        if ss_data:
            if not profile.education.k12_enrolled and ss_data.get("k12_enrollment"):
                profile.education.k12_enrolled = ss_data["k12_enrollment"]
            if not profile.education.avg_public_spend_per_student and ss_data.get("per_pupil_spending"):
                profile.education.avg_public_spend_per_student = ss_data["per_pupil_spending"]
            if not profile.education.student_teacher_ratio and ss_data.get("student_teacher_ratio"):
                profile.education.student_teacher_ratio = ss_data["student_teacher_ratio"]
            if not profile.education.teacher_count and ss_data.get(
                    "k12_enrollment") and ss_data.get("student_teacher_ratio"):
                profile.education.teacher_count = round(
                    ss_data["k12_enrollment"] / ss_data["student_teacher_ratio"]
                )

            spending_spotlight_context = (
                f"\n\n**K-12 Spending Spotlight Data (Reason Foundation, 2002-2023):**\n"
                f"Source: https://spending-spotlight.reason.org\n"
                f"- Per-pupil spending: ${ss_data.get('per_pupil_spending', 'N/A'):,} "
                f"(national avg: ${national.get('per_pupil_spending', {}).get('national_average_2023', 20322):,})\n"
                f"- Spending rank: #{ss_data.get('spending_rank', 'N/A')} nationally\n"
                f"- K-12 enrollment: {ss_data.get('k12_enrollment', 'N/A'):,}\n"
                f"- Average teacher salary: ${ss_data.get('avg_teacher_salary', 'N/A'):,} "
                f"(national avg: ${national.get('teacher_salary', {}).get('avg_2022', 70548):,})\n"
                f"- Student-teacher ratio: {ss_data.get('student_teacher_ratio', 'N/A')}\n"
                f"- Revenue per pupil: ${ss_data.get('revenue_per_pupil', 'N/A'):,}\n"
                f"- Instructional spending share: {ss_data.get('instructional_spending_pct', 'N/A')}%\n"
                f"- Benefit spending per pupil: ${ss_data.get('benefit_spending_per_pupil', 'N/A'):,}\n"
                f"- Enrollment change (2020-2023): {ss_data.get('enrollment_change_2020_2023_pct', 'N/A')}%\n"
                f"- NAEP 4th grade reading proficient: {ss_data.get('naep_4th_reading_proficient_pct', 'N/A')}%\n"
                f"- NAEP 4th grade math proficient: {ss_data.get('naep_4th_math_proficient_pct', 'N/A')}%\n"
                f"- NAEP 8th grade reading proficient: {ss_data.get('naep_8th_reading_proficient_pct', 'N/A')}%\n"
                f"- NAEP 8th grade math proficient: {ss_data.get('naep_8th_math_proficient_pct', 'N/A')}%\n"
            )
        if national:
            nat_pp = national.get('per_pupil_spending', {})
            nat_ts = national.get('teacher_salary', {})
            nat_sf = national.get('staffing', {})
            nat_en = national.get('enrollment', {})
            nat_eb = national.get('employee_benefits', {})
            nat_so = national.get('student_outcomes', {})
            spending_spotlight_context += (
                f"\n**National Benchmarks (Spending Spotlight 2025):**\n"
                f"- National avg per-pupil spending: ${nat_pp.get('national_average_2023', 20322):,}\n"
                f"- Per-pupil spending rose {nat_pp.get('change_pct', 35.8)}% (2002-2023)\n"
                f"- Avg teacher salary (national): ${nat_ts.get('avg_2022', 70548):,}\n"
                f"- Teacher salary change: {nat_ts.get('change_pct', -6.1)}% (2002-2022)\n"
                f"- Non-teaching staff growth: {nat_sf.get('non_teaching_staff_growth_pct', 22.8)}%"
                f" vs {nat_en.get('change_2002_2023_pct', 4.1)}% enrollment growth\n"
                f"- Benefit spending per pupil rose {nat_eb.get('change_pct', 81.1)}% (2002-2023)\n"
                f"- {nat_en.get('states_with_decline_2020_2023', 39)} states saw enrollment decline (2020-2023)\n"
                f"- ~{nat_so.get('naep_4th_grade_reading_below_basic_pct', 40)}%"
                f" of 4th graders below basic reading level (NAEP)\n"
            )
        if alpha_insights:
            disconnect = alpha_insights.get("spending_vs_outcomes_disconnect", {})
            state_key = target.replace(" ", "_")
            priority = alpha_insights.get("priority_state_economics", {}).get(state_key, {})
            spending_spotlight_context += (
                f"\n**Alpha Strategic Insight (from Spending Spotlight data):**\n"
                f"- Core argument: {disconnect.get('summary', 'More spending has not led to better outcomes.')}\n"
            )
            if priority:
                spending_spotlight_context += (
                    f"- Alpha intervention cost: ${priority.get('alpha_intervention_cost', 2000):,}/student "
                    f"({priority.get('alpha_pct_of_per_pupil', 'N/A')}% of per-pupil spend)\n"
                    f"- Alpha full transformation: ${priority.get('alpha_full_transform_cost', 'N/A'):,}/student "
                    f"({priority.get('alpha_full_pct_of_per_pupil', 20)}% of per-pupil spend)\n"
                    f"- Pitch: {priority.get('argument', '')}\n"
                )

    # ------------------------------------------------------------------
    # Build data context for report prompt
    # ------------------------------------------------------------------
    data_context = (
        f"**World Bank Data:**\n{wb_data}\n\n"
        f"**Live Research (Perplexity Sonar Deep Research):**\n{research_text}\n\n"
        f"**Education Research (Perplexity):**\n{edu_research_text}\n\n"
        f"**Citations:**\n{chr(10).join(str(c) for c in citations)}"
    )
    if edu_citations:
        data_context += f"\n{chr(10).join(str(c) for c in edu_citations)}"
    data_context += spending_spotlight_context

    # ------------------------------------------------------------------
    # LLM call 2: combined narrative report
    # ------------------------------------------------------------------
    if feedback and previous_report:
        report_md = await call_llm_plain(
            system_prompt=REPORT_REVISION_PROMPT.format(
                original_report=previous_report, feedback=feedback
            ),
            user_prompt=f"Revise the full briefing for {target} incorporating the feedback above.",
        )
    else:
        if target_type == TargetType.US_STATE:
            prompt = US_STATE_REPORT_PROMPT.format(target=target, data_context=data_context)
        else:
            prompt = REPORT_PROMPT.format(target=target, data_context=data_context)

        logger.info("Generating combined report for %s", target)
        report_md = await call_llm_plain(
            system_prompt=prompt,
            user_prompt=f"Write the country and education briefing for {target}.",
        )

    # ------------------------------------------------------------------
    # Save report as DOCX
    # ------------------------------------------------------------------
    docx_path = _save_report_docx(target, report_md, "Country & Education Briefing")

    logger.info("Country research complete: %s", target)
    return profile, report_md, docx_path


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _save_report_docx(target: str, markdown: str, title: str) -> str:
    """Save a markdown report as a basic DOCX file."""
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = DocxPt(11)

    heading = doc.add_heading(f"{title}: {target}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "CONFIDENTIAL & PROPRIETARY — 2hr Learning (Alpha)"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for line in markdown.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("|"):
            doc.add_paragraph(stripped, style="Normal")
        elif stripped:
            doc.add_paragraph(stripped)

    output_dir = os.path.join(OUTPUT_DIR, target.lower().replace(" ", "_"))
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"{target.replace(' ', '_')}_country_report.docx")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _merge_profiles(base: CountryProfile, synth: CountryProfile) -> None:
    """Merge synthesised profile into base, filling gaps only."""
    for model_cls, attr in [
        (Demographics, "demographics"), (Economy, "economy"),
        (EducationData, "education"), (Regulatory, "regulatory"),
        (PoliticalContext, "political_context"),
    ]:
        base_obj = getattr(base, attr)
        synth_obj = getattr(synth, attr)
        for field in model_cls.model_fields:
            if getattr(base_obj, field) is None and getattr(synth_obj, field) is not None:
                setattr(base_obj, field, getattr(synth_obj, field))

    if synth.competitive_landscape.major_operators:
        base.competitive_landscape.major_operators = synth.competitive_landscape.major_operators
    for f in ("international_chains", "edtech_penetration", "market_gaps"):
        if getattr(synth.competitive_landscape, f):
            setattr(base.competitive_landscape, f, getattr(synth.competitive_landscape, f))

    if synth.us_state_esa and base.us_state_esa is None:
        base.us_state_esa = synth.us_state_esa
    if not base.target.region and synth.target.region:
        base.target.region = synth.target.region
