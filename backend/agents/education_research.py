"""Education System Research Agent — thin wrapper (post-streamline March 16, 2026).

Education research is now MERGED into country_research.py as part of the
single combined research stage. This module is kept for backward compatibility
with pipeline imports, but delegates to country_research for the heavy lifting.

The structured EducationAnalysis is populated from the country research
structured extraction. The narrative report is a brief education-focused
excerpt from the combined country report.
"""

from __future__ import annotations

import logging
import os

from docx import Document as DocxDocument
from docx.shared import Pt as DocxPt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import (
    EducationAnalysis, SystemDiagnosis, TwoHrLearningFit, EntryMode,
    CountryProfile,
)
from services.llm import call_llm, call_llm_plain
from config import OUTPUT_DIR

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Head-of-State preamble (applied to all prompts)
# ---------------------------------------------------------------------------

HEAD_OF_STATE_PREAMBLE = """CRITICAL — EXTERNALLY PRESENTABLE ARTIFACTS:
These documents will be presented DIRECTLY to heads of state, sovereign rulers, royal family
members, and senior government ministers. The UAE version was personally reviewed by the
daughter of the #2 figure in Abu Dhabi and the #3 figure in the UAE.

ABSOLUTE RULES:
1. NEVER include internal classifications, tiers, scores, ratings, or internal frameworks
2. NEVER reference analytical methodologies, scoring matrices, or internal jargon
3. Write in the voice of a trusted advisor presenting to royalty
4. All financial figures are FIXED — research is color commentary only
5. Every artifact must be suitable for immediate external presentation
"""

# ---------------------------------------------------------------------------
# Focused education brief prompt (single call, replaces 3 x 3000-word calls)
# ---------------------------------------------------------------------------

EDUCATION_BRIEF_PROMPT = """You are a senior education advisor preparing a concise education
landscape brief for **{target}** as part of a partnership proposal with Alpha Holdings, Inc.

""" + HEAD_OF_STATE_PREAMBLE + """

## Context
{data_context}

Write a focused 2-3 page EDUCATION LANDSCAPE BRIEF covering:

## 1. Education System Snapshot
One summary table with key facts: total K-12 students, public/private split, government
education budget, per-student public spend, language of instruction, school calendar,
compulsory education ages. Then 2-3 paragraphs on the system architecture.

## 2. What's Working & What's Broken
Brief, honest assessment: outcomes (PISA/TIMSS if available), equity gaps, efficiency issues.
Focus on the 3-4 biggest pain points that Alpha's model directly addresses.
Do NOT include exhaustive scoring matrices or risk registers.

## 3. Reform Momentum & Government Appetite
Active reforms, government vision for education, appetite for innovation, precedent for
international partnerships. Name key decision-makers where known.

## 4. Localisation Considerations
Language, curriculum, cultural values, mandatory subjects, religious education requirements,
and cultural life-skills ideas (sports, arts, national interests) for the Alpha model.

IMPORTANT: This brief is COLOR COMMENTARY for the sales pitch — it does NOT drive financial
model numbers. The financial model is FIXED. Keep it focused and presentation-ready.
Do NOT include TAM/SAM calculations, competitor tables with 10+ rows, or analytical frameworks.
Maximum 1,500 words.
"""

REPORT_REVISION_PROMPT = """You are revising an education landscape brief based on feedback.

""" + HEAD_OF_STATE_PREAMBLE + """

Original brief:
{original_report}

Feedback:
{feedback}

Produce a revised brief incorporating the feedback. Keep it concise and presentation-ready.
Do NOT expand beyond 2-3 pages."""


async def run_education_research(
    target: str,
    country_profile: CountryProfile,
    feedback: str | None = None,
    previous_report: str | None = None,
) -> tuple[EducationAnalysis, str, str]:
    """Execute education research. Returns (analysis, report_md, docx_path).

    Post-streamline: this is a lightweight stage. The heavy research is done
    in country_research.py. This produces a focused education brief + structured
    analysis from the country profile data already gathered.
    """
    logger.info("Running education brief for %s", target)

    # --- Build context from existing country profile (no new Perplexity call) ---
    ctx_parts = [f"Target: {target}", f"Type: {country_profile.target.type.value}"]
    if country_profile.demographics.total_population:
        ctx_parts.append(f"Population: {country_profile.demographics.total_population:,.0f}")
    if country_profile.economy.gdp_per_capita:
        ctx_parts.append(f"GDP/cap: ${country_profile.economy.gdp_per_capita:,.0f}")
    if country_profile.education.k12_enrolled:
        ctx_parts.append(f"K-12 students: {country_profile.education.k12_enrolled:,.0f}")
    if country_profile.education.avg_public_spend_per_student:
        ctx_parts.append(f"Gov per-student spend: ${country_profile.education.avg_public_spend_per_student:,.0f}")
    if country_profile.education.avg_private_tuition:
        ctx_parts.append(f"Avg private tuition: ${country_profile.education.avg_private_tuition:,.0f}")
    if country_profile.education.literacy_rate:
        ctx_parts.append(f"Literacy rate: {country_profile.education.literacy_rate}%")
    if country_profile.education.education_budget_pct_gdp:
        ctx_parts.append(f"Education spend (% GDP): {country_profile.education.education_budget_pct_gdp}%")
    if country_profile.regulatory.foreign_ownership_rules:
        ctx_parts.append(f"Foreign ownership: {country_profile.regulatory.foreign_ownership_rules}")

    # Use any research sources from country profile
    sources_text = ""
    if country_profile.research_sources:
        sources_text = "\n**Sources from country research:**\n" + "\n".join(
            f"- {s}" for s in country_profile.research_sources[:10]
        )

    data_context = "\n".join(ctx_parts) + sources_text

    # --- Structured analysis (single call) ---
    try:
        analysis: EducationAnalysis = await call_llm(
            system_prompt=(
                HEAD_OF_STATE_PREAMBLE
                + "\nExtract structured education analysis from the country data. "
                "Focus on system diagnosis, reform landscape, and Alpha model fit."
            ),
            user_prompt=f"Context:\n{data_context}",
            output_schema=EducationAnalysis,
        )
    except Exception as exc:
        logger.warning("Structured analysis failed: %s", exc)
        analysis = EducationAnalysis(
            system_diagnosis=SystemDiagnosis(primary_pain_points=["Analysis pending"]),
            two_hr_learning_fit=TwoHrLearningFit(model_recommendation=EntryMode.HYBRID),
        )

    # --- Narrative brief (single call) ---
    if feedback and previous_report:
        report_md = await call_llm_plain(
            system_prompt=REPORT_REVISION_PROMPT.format(
                original_report=previous_report, feedback=feedback
            ),
            user_prompt=f"Revise the education brief for {target}.",
        )
    else:
        report_md = await call_llm_plain(
            system_prompt=EDUCATION_BRIEF_PROMPT.format(
                target=target, data_context=data_context
            ),
            user_prompt=f"Write a concise education landscape brief for {target}.",
        )

    # --- Save as DOCX ---
    docx_path = _save_report_docx(target, report_md, "Education Landscape Brief")

    logger.info("Education brief complete for %s", target)
    return analysis, report_md, docx_path


def _save_report_docx(target: str, markdown: str, title: str) -> str:
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = DocxPt(11)
    h = doc.add_heading(f"{title}: {target}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("CONFIDENTIAL & PROPRIETARY — Alpha Holdings, Inc.").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    for line in markdown.split("\n"):
        s = line.strip()
        if s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s.startswith("|"):
            doc.add_paragraph(s)
        elif s:
            doc.add_paragraph(s)
    output_dir = os.path.join(OUTPUT_DIR, target.lower().replace(" ", "_"))
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"{target.replace(' ', '_')}_education_brief.docx")
    doc.save(path)
    return path
